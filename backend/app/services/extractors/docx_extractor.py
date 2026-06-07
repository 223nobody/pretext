from __future__ import annotations

from pathlib import Path


async def extract_docx(file_path: Path, _encoding: str) -> tuple[str, dict]:
    from docx import Document

    document = Document(str(file_path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n\n".join(parts), {}
